import logging
import os
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile
import shutil
from datetime import datetime
from openai import AsyncOpenAI
from .schemas import FileUpload, ConversationContext
from .config import config

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles file upload, processing, and context-aware analysis"""
    
    def __init__(self):
        self.llm = AsyncOpenAI(
            api_key=config.openai_api_key
        )
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        self.supported_file_types = config.supported_file_types
    
    async def process_uploaded_file(
        self, 
        file_data: bytes, 
        filename: str, 
        context: ConversationContext
    ) -> FileUpload:
        """Process an uploaded file and create metadata"""
        
        try:
            # Generate unique file ID
            file_id = self._generate_file_id(file_data, filename)
            
            # Validate file type and size
            file_type = self._get_file_type(filename)
            file_size = len(file_data)
            
            if not self._is_file_type_supported(file_type):
                raise ValueError(f"Unsupported file type: {file_type}")
            
            if file_size > config.max_file_size:
                raise ValueError(f"File too large: {file_size} bytes (max: {config.max_file_size})")
            
            # Save file to upload directory
            file_path = self.upload_dir / f"{file_id}_{filename}"
            with open(file_path, "wb") as f:
                f.write(file_data)
            
            # Extract content and metadata
            content = await self._extract_file_content(file_path, file_type)
            metadata = await self._extract_file_metadata(file_path, file_type, content)
            
            # Create file upload record
            file_upload = FileUpload(
                file_id=file_id,
                filename=filename,
                file_type=file_type,
                file_size=file_size,
                upload_time=datetime.utcnow(),
                content_hash=hashlib.md5(file_data).hexdigest()
            )
            
            # Store content and metadata for context
            await self._store_file_context(file_id, content, metadata, context)
            
            logger.info(f"Successfully processed file: {filename} (ID: {file_id})")
            return file_upload
            
        except Exception as e:
            logger.error(f"Error processing uploaded file {filename}: {e}")
            raise
    
    def _generate_file_id(self, file_data: bytes, filename: str) -> str:
        """Generate unique file ID based on content and filename"""
        
        # Combine file content hash with filename hash
        content_hash = hashlib.md5(file_data).hexdigest()[:8]
        filename_hash = hashlib.md5(filename.encode()).hexdigest()[:8]
        timestamp = str(int(datetime.utcnow().timestamp()))[-6:]
        
        return f"{content_hash}_{filename_hash}_{timestamp}"
    
    def _get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        
        # Get MIME type
        mime_type, _ = mimetypes.guess_type(filename)
        
        if mime_type:
            return mime_type
        
        # Fallback to extension-based detection
        extension = Path(filename).suffix.lower()
        
        extension_map = {
            ".pdf": "application/pdf",
            ".csv": "text/csv",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/msword",
            ".txt": "text/plain"
        }
        
        return extension_map.get(extension, "application/octet-stream")
    
    def _is_file_type_supported(self, file_type: str) -> bool:
        """Check if file type is supported"""
        
        supported_types = [
            "application/pdf",
            "text/csv",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain"
        ]
        
        return file_type in supported_types
    
    async def _extract_file_content(self, file_path: Path, file_type: str) -> str:
        """Extract text content from file based on type"""
        
        try:
            if file_type == "text/csv":
                return await self._extract_csv_content(file_path)
            elif file_type == "text/plain":
                return await self._extract_text_content(file_path)
            elif "spreadsheet" in file_type or "excel" in file_type:
                return await self._extract_excel_content(file_path)
            elif "word" in file_type or "document" in file_type:
                return await self._extract_word_content(file_path)
            elif file_type == "application/pdf":
                return await self._extract_pdf_content(file_path)
            else:
                return await self._extract_generic_content(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return f"Error extracting content: {str(e)}"
    
    async def _extract_csv_content(self, file_path: Path) -> str:
        """Extract content from CSV file"""
        
        try:
            import pandas as pd
            
            # Read CSV with pandas
            df = pd.read_csv(file_path)
            
            # Convert to readable format
            content = f"CSV File: {file_path.name}\n"
            content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n\n"
            content += "Column Headers:\n"
            content += ", ".join(df.columns.tolist()) + "\n\n"
            
            # Add first few rows as sample
            content += "Sample Data:\n"
            content += df.head().to_string(index=False)
            
            return content
            
        except ImportError:
            logger.warning("pandas not available, using basic CSV parsing")
            return await self._extract_text_content(file_path)
        except Exception as e:
            logger.error(f"Error parsing CSV with pandas: {e}")
            return await self._extract_text_content(file_path)
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """Extract content from text file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Limit content length for processing
            if len(content) > 10000:
                content = content[:10000] + "\n\n[Content truncated for processing]"
            
            return content
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    if len(content) > 10000:
                        content = content[:10000] + "\n\n[Content truncated for processing]"
                    
                    return content
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, return error message
            return f"Error: Could not decode file content with any supported encoding"
    
    async def _extract_excel_content(self, file_path: Path) -> str:
        """Extract content from Excel file"""
        
        try:
            import pandas as pd
            
            # Read Excel file
            excel_file = pd.ExcelFile(file_path)
            content = f"Excel File: {file_path.name}\n"
            content += f"Sheets: {excel_file.sheet_names}\n\n"
            
            # Process each sheet
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    content += f"Sheet: {sheet_name}\n"
                    content += f"Rows: {len(df)}, Columns: {len(df.columns)}\n"
                    content += "Columns: " + ", ".join(df.columns.tolist()) + "\n"
                    content += "Sample Data:\n"
                    content += df.head().to_string(index=False) + "\n\n"
                except Exception as e:
                    content += f"Error reading sheet {sheet_name}: {str(e)}\n\n"
            
            # Limit content length
            if len(content) > 10000:
                content = content[:10000] + "\n\n[Content truncated for processing]"
            
            return content
            
        except ImportError:
            logger.warning("pandas not available, cannot read Excel file")
            return f"Error: Excel file processing requires pandas library"
        except Exception as e:
            logger.error(f"Error parsing Excel file: {e}")
            return f"Error parsing Excel file: {str(e)}"
    
    async def _extract_word_content(self, file_path: Path) -> str:
        """Extract content from Word document"""
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = f"Word Document: {file_path.name}\n\n"
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                content += "\n[Table Content]\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
            
            # Limit content length
            if len(content) > 10000:
                content = content[:10000] + "\n\n[Content truncated for processing]"
            
            return content
            
        except ImportError:
            logger.warning("python-docx not available, cannot read Word document")
            return f"Error: Word document processing requires python-docx library"
        except Exception as e:
            logger.error(f"Error parsing Word document: {e}")
            return f"Error parsing Word document: {str(e)}"
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF file"""
        
        try:
            import PyPDF2
            
            content = f"PDF File: {file_path.name}\n\n"
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                content += f"Pages: {len(pdf_reader.pages)}\n\n"
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    if page_num >= 5:  # Limit to first 5 pages
                        content += f"\n[Additional pages truncated for processing]\n"
                        break
                    
                    page_text = page.extract_text()
                    if page_text.strip():
                        content += f"Page {page_num + 1}:\n{page_text}\n\n"
            
            # Limit content length
            if len(content) > 10000:
                content = content[:10000] + "\n\n[Content truncated for processing]"
            
            return content
            
        except ImportError:
            logger.warning("PyPDF2 not available, cannot read PDF file")
            return f"Error: PDF processing requires PyPDF2 library"
        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            return f"Error parsing PDF file: {str(e)}"
    
    async def _extract_generic_content(self, file_path: Path) -> str:
        """Extract content from generic file types"""
        
        try:
            # Try to read as text first
            return await self._extract_text_content(file_path)
        except Exception:
            return f"File type not supported for content extraction: {file_path.name}"
    
    async def _extract_file_metadata(self, file_path: Path, file_type: str, content: str) -> Dict[str, Any]:
        """Extract metadata from file"""
        
        try:
            metadata = {
                "file_type": file_type,
                "file_size": file_path.stat().st_size,
                "created_time": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                "content_length": len(content),
                "content_preview": content[:500] + "..." if len(content) > 500 else content
            }
            
            # Add file-specific metadata
            if file_type == "text/csv":
                metadata.update(await self._extract_csv_metadata(file_path, content))
            elif "spreadsheet" in file_type or "excel" in file_type:
                metadata.update(await self._extract_excel_metadata(file_path, content))
            elif "word" in file_type or "document" in file_type:
                metadata.update(await self._extract_word_metadata(file_path, content))
            elif file_type == "application/pdf":
                metadata.update(await self._extract_pdf_metadata(file_path, content))
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting file metadata: {e}")
            return {"error": str(e)}
    
    async def _extract_csv_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract CSV-specific metadata"""
        
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            return {
                "csv_rows": len(df),
                "csv_columns": len(df.columns),
                "csv_column_names": df.columns.tolist(),
                "csv_data_types": df.dtypes.to_dict()
            }
        except Exception:
            return {}
    
    async def _extract_excel_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract Excel-specific metadata"""
        
        try:
            import pandas as pd
            
            excel_file = pd.ExcelFile(file_path)
            
            return {
                "excel_sheets": excel_file.sheet_names,
                "excel_sheet_count": len(excel_file.sheet_names)
            }
        except Exception:
            return {}
    
    async def _extract_word_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract Word-specific metadata"""
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            return {
                "word_paragraphs": len(doc.paragraphs),
                "word_tables": len(doc.tables),
                "word_sections": len(doc.sections)
            }
        except Exception:
            return {}
    
    async def _extract_pdf_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                return {
                    "pdf_pages": len(pdf_reader.pages),
                    "pdf_info": pdf_reader.metadata if pdf_reader.metadata else {}
                }
        except Exception:
            return {}
    
    async def _store_file_context(
        self, 
        file_id: str, 
        content: str, 
        metadata: Dict[str, Any], 
        context: ConversationContext
    ):
        """Store file context for conversation use"""
        
        try:
            # Store in temporary context (in production, use Redis or database)
            if not hasattr(self, '_file_contexts'):
                self._file_contexts = {}
            
            self._file_contexts[file_id] = {
                "content": content,
                "metadata": metadata,
                "context": context,
                "timestamp": datetime.utcnow()
            }
            
            logger.info(f"Stored file context for {file_id}")
            
        except Exception as e:
            logger.error(f"Error storing file context: {e}")
    
    async def analyze_file_in_context(
        self, 
        file_id: str, 
        query: str, 
        context: ConversationContext
    ) -> str:
        """Analyze uploaded file in the context of conversation"""
        
        try:
            # Get file context
            if not hasattr(self, '_file_contexts') or file_id not in self._file_contexts:
                return f"Error: File {file_id} not found in context"
            
            file_context = self._file_contexts[file_id]
            file_content = file_context["content"]
            file_metadata = file_context["metadata"]
            
            # Build analysis prompt
            prompt = self._build_file_analysis_prompt(query, file_content, file_metadata, context)
            
            # Get LLM analysis
            response = await self.llm.chat.completions.create(
                model=config.openai_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1
            )
            
            analysis = response.choices[0].message.content
            
            logger.info(f"Generated file analysis for {file_id}")
            return analysis
            
        except Exception as e:
            logger.error(f"Error analyzing file {file_id}: {e}")
            return f"Error analyzing file: {str(e)}"
    
    def _build_file_analysis_prompt(
        self, 
        query: str, 
        file_content: str, 
        file_metadata: Dict[str, Any], 
        context: ConversationContext
    ) -> str:
        """Build prompt for file analysis"""
        
        return f"""
        Analyze this uploaded file in the context of the user's question and conversation history.
        
        **User Question:** "{query}"
        
        **Conversation Context:**
        - User's knowledge level: {context.knowledge_level.value}
        - Previous questions: {', '.join(context.semantic_themes) if context.semantic_themes else 'None'}
        - Current focus: {context.current_topic or 'General'}
        
        **File Information:**
        - Filename: {file_metadata.get('filename', 'Unknown')}
        - File type: {file_metadata.get('file_type', 'Unknown')}
        - File size: {file_metadata.get('file_size', 'Unknown')} bytes
        - Content length: {file_metadata.get('content_length', 'Unknown')} characters
        
        **File Content:**
        {file_content}
        
        **Analysis Requirements:**
        1. **Answer the user's question** using information from the file
        2. **Provide insights** relevant to their financial planning needs
        3. **Extract key data points** that are relevant to their situation
        4. **Suggest next steps** based on the file content
        5. **Identify any limitations** or areas needing clarification
        
        **Response Guidelines:**
        - Be specific and actionable
        - Reference specific parts of the file when possible
        - Consider the user's knowledge level
        - Focus on financial planning and life insurance context
        - Provide educational value, not specific financial advice
        
        **Response Structure:**
        - Direct answer to their question
        - Key insights from the file
        - Relevant data points
        - Recommendations or next steps
        - Any questions for clarification
        
        Generate a comprehensive analysis:
        """
    
    def get_file_summary(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Get summary of uploaded file"""
        
        try:
            if hasattr(self, '_file_contexts') and file_id in self._file_contexts:
                file_context = self._file_contexts[file_id]
                return {
                    "file_id": file_id,
                    "filename": file_context["metadata"].get("filename", "Unknown"),
                    "file_type": file_context["metadata"].get("file_type", "Unknown"),
                    "upload_time": file_context["timestamp"].isoformat(),
                    "content_preview": file_context["metadata"].get("content_preview", "")[:200] + "..."
                }
            return None
            
        except Exception as e:
            logger.error(f"Error getting file summary: {e}")
            return None
    
    def cleanup_file(self, file_id: str) -> bool:
        """Clean up file and context"""
        
        try:
            # Remove file context
            if hasattr(self, '_file_contexts') and file_id in self._file_contexts:
                del self._file_contexts[file_id]
            
            # Remove physical file
            for file_path in self.upload_dir.glob(f"{file_id}_*"):
                file_path.unlink()
                logger.info(f"Removed file: {file_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"Error cleaning up file {file_id}: {e}")
            return False
    
    def cleanup_old_files(self, max_age_hours: int = 24) -> int:
        """Clean up old uploaded files"""
        
        try:
            cleaned_count = 0
            current_time = datetime.utcnow()
            
            # Clean up file contexts
            if hasattr(self, '_file_contexts'):
                file_ids_to_remove = []
                
                for file_id, context in self._file_contexts.items():
                    age_hours = (current_time - context["timestamp"]).total_seconds() / 3600
                    if age_hours > max_age_hours:
                        file_ids_to_remove.append(file_id)
                
                for file_id in file_ids_to_remove:
                    del self._file_contexts[file_id]
                    cleaned_count += 1
            
            # Clean up physical files
            for file_path in self.upload_dir.glob("*"):
                if file_path.is_file():
                    file_age_hours = (current_time - datetime.fromtimestamp(file_path.stat().st_mtime)).total_seconds() / 3600
                    if file_age_hours > max_age_hours:
                        file_path.unlink()
                        cleaned_count += 1
            
            logger.info(f"Cleaned up {cleaned_count} old files")
            return cleaned_count
            
        except Exception as e:
            logger.error(f"Error cleaning up old files: {e}")
            return 0 